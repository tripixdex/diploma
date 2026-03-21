#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import sys
import argparse
from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_section_margins(doc: Document):
    for sec in doc.sections:
        sec.left_margin   = Mm(30)
        sec.right_margin  = Mm(10)
        sec.top_margin    = Mm(20)
        sec.bottom_margin = Mm(20)

def force_normal_style(doc: Document):
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(14)
    pf = st.paragraph_format
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(1.25)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

def remove_paragraph(paragraph):
    el = paragraph._element
    el.getparent().remove(el)

def is_toc_paragraph(p) -> bool:
    n = (p.style.name or "").lower() if p.style else ""
    return n.startswith("toc")

def remove_toc_self_entry(doc: Document):
    # удаляем строку оглавления "СОДЕРЖАНИЕ\t2" (самоссылка), если она есть
    for p in list(doc.paragraphs):
        txt = (p.text or "").strip()
        if is_toc_paragraph(p) and txt.startswith("СОДЕРЖАНИЕ") and "\t" in txt:
            left, right = txt.rsplit("\t", 1)
            if left.strip() == "СОДЕРЖАНИЕ" and right.strip().isdigit():
                remove_paragraph(p)
                break

def set_page_number_start(doc: Document, start: int):
    """
    В sectPr добавляет/обновляет <w:pgNumType w:start="N"/>.
    Важно: python-docx BaseOxmlElement.xpath НЕ принимает namespaces=..., поэтому используем xpath без него.
    """
    sectPr = doc.sections[0]._sectPr
    found = sectPr.xpath("./w:pgNumType")
    if found:
        pg = found[0]
    else:
        pg = OxmlElement("w:pgNumType")
        sectPr.append(pg)
    pg.set(qn("w:start"), str(start))

def normalize_paragraphs(doc: Document):
    # Нормализуем только "обычные" параграфы: без “конских” интервалов.
    for p in doc.paragraphs:
        # шрифт в ран-ах (в т.ч. в заголовках)
        for r in p.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(14)

        st = (p.style.name or "").lower() if p.style else ""
        pf = p.paragraph_format

        if st.startswith("heading") or "заголовок" in st:
            # умеренные интервалы, чтобы не было провалов
            pf.space_before = Pt(12)
            pf.space_after  = Pt(6)
            pf.first_line_indent = None
            # заголовки обычно слева по методичке/принятой практике кафедры
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf.line_spacing = 1.5
            pf.first_line_indent = Cm(1.25)
            pf.space_before = Pt(0)
            pf.space_after  = Pt(0)

def remove_empty_paragraphs(doc: Document):
    # Удаляем пустые параграфы, которые создают “пустые страницы/провалы”.
    for p in list(doc.paragraphs):
        if (p.text or "").strip() == "" and len(p.runs) == 0:
            remove_paragraph(p)

def normalize_tables(doc: Document):
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = "Times New Roman"
                        r.font.size = Pt(14)

def main(inp: str, out: str, start_page: int):
    doc = Document(inp)
    set_section_margins(doc)
    force_normal_style(doc)

    # опционально: старт нумерации (2 если перед этим будет титул как стр.1 без номера)
    if start_page is not None:
        set_page_number_start(doc, start=start_page)

    remove_toc_self_entry(doc)
    normalize_paragraphs(doc)
    remove_empty_paragraphs(doc)
    normalize_tables(doc)

    doc.save(out)
    print(f"OK: {out}")

if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("input_docx")
    ap.add_argument("output_docx")
    ap.add_argument("--start-page", type=int, default=2,
                    help="Старт нумерации страниц (2 если титул отдельным файлом как стр.1 без номера).")
    args = ap.parse_args()
    main(args.input_docx, args.output_docx, start_page=args.start_page)
