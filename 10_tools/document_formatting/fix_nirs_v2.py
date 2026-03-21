#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
fix_nirs_v2.py — минимальный авто-фикс оформления НИРС/РПЗ под методичку.
Исправляет:
  - разрывы страниц перед каждым заголовком Heading 1 (кроме первого)
  - интервалы после заголовков (3 межстрочных интервала; 2 если сразу следующий заголовок)
  - оформление подписей таблиц (влево, без абзацного отступа)
  - перенумерацию таблиц (по порядку появления, начиная с 1)
  - “псевдо-списки” в Normal со строками, начинающимися на "—" (переводит в List Bullet)
  - вставку страницы "ПРИЛОЖЕНИЯ" перед первым приложением и "Приложение А" вместо "ПРИЛОЖЕНИЕ А"

Требования:
  pip install python-docx

Пример:
  python fix_nirs_v2.py in.docx out.docx
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

EMU_PER_INCH = 914400

def emu_to_mm(emu: int) -> float:
    return emu / EMU_PER_INCH * 25.4

def pt_to_emu(pt: float) -> int:
    return int(pt * 12700)

def para_has_page_break(p) -> bool:
    # page break in runs or pageBreakBefore in pPr
    if p._p.xpath('.//w:br[@w:type="page"]'):
        return True
    if p._p.xpath('./w:pPr/w:pageBreakBefore'):
        return True
    return False

def fix_doc(doc: Document) -> None:
    # ВАЖНО: этот фикс предполагает, что в данном файле НЕТ титульного листа.
    # Поэтому нумерация страниц должна быть видимой уже на странице "СОДЕРЖАНИЕ".
    doc.sections[0].different_first_page_header_footer = False

    # 1) Разрывы страниц перед каждым Heading 1 (кроме первого)
    first_h1 = True
    for p in doc.paragraphs:
        if p.style.name == "Heading 1" and p.text.strip():
            if first_h1:
                first_h1 = False
                continue
            if not para_has_page_break(p):
                p.paragraph_format.page_break_before = True

    # 2) Интервалы после заголовков: 3 интервала (≈63 pt), если далее текст; 2 (≈42 pt), если далее заголовок
    after_text = pt_to_emu(63)     # 3 * (14pt * 1.5) = 63pt
    after_heading = pt_to_emu(42)  # 2 * (14pt * 1.5) = 42pt

    paras = doc.paragraphs
    for i, p in enumerate(paras):
        if p.style.name in ("Heading 1", "Heading 2", "Heading 3") and p.text.strip():
            j = i + 1
            while j < len(paras) and not paras[j].text.strip():
                j += 1
            if j < len(paras) and paras[j].style.name.startswith("Heading"):
                p.paragraph_format.space_after = after_heading
            else:
                p.paragraph_format.space_after = after_text
            p.paragraph_format.space_before = 0

    # 3) “псевдо-списки”: Normal + левый отступ 1.25см + строка начинается с "—"
    for p in doc.paragraphs:
        if p.style.name != "Normal":
            continue
        txt = p.text.strip()
        fmt = p.paragraph_format
        if not txt.startswith("—"):
            continue
        if fmt.left_indent is None:
            continue
        if abs(emu_to_mm(fmt.left_indent) - 12.5) > 0.9:
            continue

        try:
            p.style = doc.styles["List Bullet"]
        except KeyError:
            # если в документе нет List Bullet — оставляем как есть
            continue
        p.text = re.sub(r'^—\s*', '', txt)
        fmt.left_indent = None
        fmt.first_line_indent = None

    # 4) Подписи таблиц: выравнивание влево, без отступа + перенумерация
    captions = []
    for p in doc.paragraphs:
        m = re.match(r'^Таблица\s+(\d+)\s+—\s*(.*)$', p.text.strip())
        if m:
            captions.append((p, m.group(2)))

    for n, (p, title) in enumerate(captions, start=1):
        p.text = f"Таблица {n} — {title}".strip()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.first_line_indent = 0
        pf.left_indent = 0
        pf.right_indent = 0

    # 5) Приложения: вставка страницы "ПРИЛОЖЕНИЯ" перед первым приложением
    for p in doc.paragraphs:
        if p.style.name == "Heading 1" and p.text.strip().upper().startswith("ПРИЛОЖЕНИЕ"):
            new_p = p.insert_paragraph_before("ПРИЛОЖЕНИЯ")
            new_p.style = doc.styles["Heading 1"]
            new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            new_p.paragraph_format.page_break_before = True

            m = re.match(r'^ПРИЛОЖЕНИЕ\s+([А-ЯA-Z0-9]+)$', p.text.strip().upper())
            if m:
                p.text = f"Приложение {m.group(1)}"
            break

def main() -> int:
    if len(sys.argv) != 3:
        print("Usage: python fix_nirs_v2.py <in.docx> <out.docx>", file=sys.stderr)
        return 2
    in_path = Path(sys.argv[1]).expanduser()
    out_path = Path(sys.argv[2]).expanduser()

    if not in_path.exists():
        print(f"Input not found: {in_path}", file=sys.stderr)
        return 1

    doc = Document(str(in_path))
    fix_doc(doc)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Saved: {out_path}")
    return 0

if __name__ == "__main__":
    raise SystemExit(main())
