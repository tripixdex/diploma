#!/usr/bin/env python3
import re
import sys
from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None

def set_section_layout(doc):
    for s in doc.sections:
        s.page_width = Mm(210)
        s.page_height = Mm(297)
        s.left_margin = Mm(30)
        s.right_margin = Mm(10)
        s.top_margin = Mm(20)
        s.bottom_margin = Mm(20)

def set_styles(doc):
    # Normal (основной текст)
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(14)
    pf = st.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(1.25)
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    # Heading 1 (структурные элементы: ВВЕДЕНИЕ, СОДЕРЖАНИЕ, ...)
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Times New Roman"
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1pf = h1.paragraph_format
    h1pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1pf.space_before = Pt(0)
    h1pf.space_after = Pt(12)  # визуально ~ одна пустая строка

    # Heading 2 (разделы/подразделы)
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Times New Roman"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2pf = h2.paragraph_format
    h2pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    h2pf.first_line_indent = Cm(1.25)
    h2pf.space_before = Pt(0)
    h2pf.space_after = Pt(12)

def find_body_start(doc):
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().upper() == "ВВЕДЕНИЕ" and "\t" not in p.text:
            return i
    return 0

def clear_heading_run_overrides(doc):
    # убираем ручные 12pt в заголовках, чтобы работал стиль
    for p in doc.paragraphs:
        if p.style.name in ("Heading 1", "Heading 2"):
            for r in p.runs:
                if r.font.size is not None and abs(r.font.size.pt - 12.0) < 0.01:
                    r.font.size = None

def fix_broken_refs(doc, body_start):
    # заменяем битые ссылки на рисунки приложений на нормальные
    for p in doc.paragraphs[body_start:]:
        t = p.text
        t2 = t.replace("\\1А.2", "рисунке А.2").replace("\\1А.3", "рисунке А.3")
        if t2 != t:
            p.text = t2

def normalize_body_paragraphs(doc, body_start):
    for p in doc.paragraphs[body_start:]:
        if p.style.name == "Normal":
            pf = p.paragraph_format
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf.first_line_indent = Cm(1.25)
            pf.line_spacing = 1.5
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)

def cleanup_empty_lines_after_headings(doc, body_start):
    i = body_start
    while i < len(doc.paragraphs) - 1:
        p = doc.paragraphs[i]
        if p.style.name in ("Heading 1", "Heading 2"):
            # удаляем пустые параграфы сразу после заголовка
            j = i + 1
            while j < len(doc.paragraphs):
                nxt = doc.paragraphs[j]
                if nxt.text.strip() == "" and nxt.style.name == "Normal":
                    delete_paragraph(nxt)
                else:
                    break
            # i не увеличиваем по j, т.к. список изменился
        i += 1

def fix_title_page_artifact(doc):
    # строка "Таблица 1 — Тема: ..." на титуле: превращаем в обычную "Тема: ..."
    for p in doc.paragraphs[:20]:
        if p.text.strip().startswith("Таблица 1 — Тема:"):
            p.text = p.text.replace("Таблица 1 — ", "")
            break

def main(inp, outp):
    doc = Document(inp)

    set_section_layout(doc)
    set_styles(doc)

    body_start = find_body_start(doc)

    fix_title_page_artifact(doc)
    clear_heading_run_overrides(doc)
    fix_broken_refs(doc, body_start)
    normalize_body_paragraphs(doc, body_start)
    cleanup_empty_lines_after_headings(doc, body_start)

    doc.save(outp)

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: fix_nirs.py IN.docx OUT.docx", file=sys.stderr)
        sys.exit(2)
    main(sys.argv[1], sys.argv[2])
