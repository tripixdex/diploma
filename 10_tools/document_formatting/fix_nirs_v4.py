#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
fix_nirs_v4.py
Правки оформления НИРС по "Методика ВКР_Б - 2019" (РК9):
- поля: слева 30 мм, справа 10 мм, сверху/снизу 20 мм
- основной текст: Times New Roman 14, межстрочный 1.5, абзац 1.25 см, выравнивание по ширине
- заголовки: убираем "конские" интервалы (ошибка прошлых версий — space_after=63 pt)
- удаляем лишние ручные разрывы страниц, создающие пустые страницы (double page-break)

Ограничение: python-docx НЕ пересчитывает/не обновляет автоматическое оглавление (TOC).
После запуска в Word: ПКМ по оглавлению -> "Обновить поле" -> "Обновить целиком".
"""

import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def delete_paragraph(paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)


def fix_document(in_path: Path, out_path: Path) -> None:
    doc = Document(str(in_path))

    # 1) Поля страницы (по методичке)
    for sec in doc.sections:
        sec.top_margin = Mm(20)
        sec.bottom_margin = Mm(20)
        sec.left_margin = Mm(30)
        sec.right_margin = Mm(10)

    # 2) Основной стиль текста
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)

    npf = normal.paragraph_format
    npf.line_spacing = 1.5
    npf.first_line_indent = Cm(1.25)
    npf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    npf.space_before = Pt(0)
    npf.space_after = Pt(0)

    # 3) Нормируем интервалы после заголовков (ключевая правка v4)
    H1_AFTER = Pt(12)  # ~1 строка
    H2_AFTER = Pt(6)   # меньше, чем H1; чтобы не было "пустых полотен"
    H3_AFTER = Pt(0)

    for p in doc.paragraphs:
        st = p.style.name
        pf = p.paragraph_format

        if st == "Heading 1":
            pf.space_before = Pt(0)
            pf.space_after = H1_AFTER
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif st == "Heading 2":
            pf.space_before = Pt(0)
            pf.space_after = H2_AFTER
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.first_line_indent = Cm(0)

        elif st == "Heading 3":
            pf.space_before = Pt(0)
            pf.space_after = H3_AFTER
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.first_line_indent = Cm(0)

        elif st == "Normal":
            # Подстраховка: если где-то формат "переехал" и не наследуется от Normal
            if pf.first_line_indent is None:
                pf.first_line_indent = Cm(1.25)
            if pf.alignment is None:
                pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if pf.line_spacing is None:
                pf.line_spacing = 1.5
            if pf.space_before is None:
                pf.space_before = Pt(0)
            if pf.space_after is None:
                pf.space_after = Pt(0)

    # 4) Удаляем "ручные" page-break-параграфы (пустые) перед заголовками с pageBreakBefore=True,
    #    чтобы не было пустых страниц (double page-break).
    i = 0
    while i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        is_empty = (p.text.strip() == "")
        has_page_break = ('w:type="page"' in p._p.xml)

        if is_empty and has_page_break and i + 1 < len(doc.paragraphs):
            nxt = doc.paragraphs[i + 1]
            if nxt.paragraph_format.page_break_before:
                delete_paragraph(p)
                continue
        i += 1

    # 5) Убираем лишние пустые строки (оставляем максимум одну подряд)
    i = 0
    while i < len(doc.paragraphs) - 1:
        p = doc.paragraphs[i]
        nxt = doc.paragraphs[i + 1]
        if (
            p.text.strip() == ""
            and nxt.text.strip() == ""
            and p.style.name == "Normal"
            and nxt.style.name == "Normal"
        ):
            delete_paragraph(nxt)
            continue
        i += 1

    doc.save(str(out_path))


def main(argv: list[str]) -> int:
    if len(argv) != 3:
        print("USAGE:\n  python fix_nirs_v4.py <input.docx> <output.docx>", file=sys.stderr)
        return 2

    in_path = Path(argv[1]).expanduser().resolve()
    out_path = Path(argv[2]).expanduser().resolve()

    if not in_path.exists():
        print(f"ERROR: input not found: {in_path}", file=sys.stderr)
        return 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    fix_document(in_path, out_path)
    print(f"OK: saved -> {out_path}")
    print("NOTE: обнови оглавление в Word: ПКМ по оглавлению -> Обновить поле -> Обновить целиком.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
